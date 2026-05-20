"""Tests for POST /upload-template endpoint."""

import io
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest
from fastapi.testclient import TestClient

from app.main import app


@pytest.fixture
def client() -> TestClient:
    """Create test client."""
    return TestClient(app)


@pytest.fixture
def tmp_template_dir(tmp_path: Path):
    """Override template dir to use a temporary directory."""
    with patch("app.routes.upload.settings") as mock_settings:
        mock_settings.pdf_template_dir = tmp_path
        yield tmp_path


def _create_minimal_docx() -> bytes:
    """Create a minimal valid .docx file using python-docx."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello {{ nama_alat }}")
    doc.add_paragraph("{% for item in sensors %}{{ item.name }}{% endfor %}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _create_empty_docx() -> bytes:
    """Create a minimal .docx without variables."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Simple document without variables")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestUploadValidation:
    """Tests for upload validation logic."""

    def test_rejects_non_docx_file(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should return 400 for non-.docx files."""
        response = client.post(
            "/upload-template",
            files={"file": ("test.pdf", b"fake content", "application/pdf")},
            data={"template_id": "test-123", "section": "cover"},
        )
        assert response.status_code == 400
        data = response.json()
        assert "Format file tidak didukung" in data["detail"]

    def test_rejects_file_over_10mb(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should return 400 for files exceeding 10MB."""
        large_content = b"x" * (10 * 1024 * 1024 + 1)
        response = client.post(
            "/upload-template",
            files={"file": ("test.docx", large_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "test-123", "section": "cover"},
        )
        assert response.status_code == 400
        data = response.json()
        assert "Ukuran file melebihi batas maksimal 10MB" in data["detail"]

    def test_rejects_invalid_template_id_special_chars(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should return 400 for template_id with special characters."""
        docx_content = _create_empty_docx()
        response = client.post(
            "/upload-template",
            files={"file": ("test.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "test@invalid!", "section": "cover"},
        )
        assert response.status_code == 400
        data = response.json()
        assert "template_id" in data["detail"].lower() or "template_id" in data.get("field", "")

    def test_rejects_template_id_over_100_chars(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should return 400 for template_id exceeding 100 characters."""
        docx_content = _create_empty_docx()
        long_id = "a" * 101
        response = client.post(
            "/upload-template",
            files={"file": ("test.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": long_id, "section": "cover"},
        )
        assert response.status_code == 400

    def test_rejects_empty_template_id(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should return 400 for empty template_id."""
        docx_content = _create_empty_docx()
        response = client.post(
            "/upload-template",
            files={"file": ("test.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "", "section": "cover"},
        )
        assert response.status_code == 400

    def test_rejects_invalid_section(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should return 422 for invalid section value (FastAPI enum validation)."""
        docx_content = _create_empty_docx()
        response = client.post(
            "/upload-template",
            files={"file": ("test.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "test-123", "section": "invalid"},
        )
        assert response.status_code == 422


class TestUploadSuccess:
    """Tests for successful upload scenarios."""

    def test_upload_saves_file_correctly(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should save file to correct path and return 201."""
        docx_content = _create_minimal_docx()
        response = client.post(
            "/upload-template",
            files={"file": ("template.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "test-123", "section": "cover"},
        )
        assert response.status_code == 201
        data = response.json()

        # Verify response fields
        assert data["path"] == "test-123/cover.docx"
        assert data["size_bytes"] == len(docx_content)
        assert isinstance(data["variables"], list)
        assert isinstance(data["loops"], list)

        # Verify file was saved
        saved_file = tmp_template_dir / "test-123" / "cover.docx"
        assert saved_file.exists()
        assert saved_file.read_bytes() == docx_content

    def test_upload_detects_variables(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should detect variables and loops in uploaded .docx."""
        docx_content = _create_minimal_docx()
        response = client.post(
            "/upload-template",
            files={"file": ("template.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "tmpl-abc", "section": "results"},
        )
        assert response.status_code == 201
        data = response.json()

        assert "nama_alat" in data["variables"]
        assert "sensors" in data["loops"]

    def test_upload_overwrites_existing_file(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should overwrite existing file at the same path."""
        docx_content_1 = _create_empty_docx()
        docx_content_2 = _create_minimal_docx()

        # First upload
        client.post(
            "/upload-template",
            files={"file": ("template.docx", docx_content_1, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "overwrite-test", "section": "cover"},
        )

        # Second upload (overwrite)
        response = client.post(
            "/upload-template",
            files={"file": ("template.docx", docx_content_2, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "overwrite-test", "section": "cover"},
        )
        assert response.status_code == 201

        # Verify file was overwritten
        saved_file = tmp_template_dir / "overwrite-test" / "cover.docx"
        assert saved_file.read_bytes() == docx_content_2

    def test_upload_accepts_valid_template_ids(self, client: TestClient, tmp_template_dir: Path) -> None:
        """Should accept template_id with alphanumeric, hyphen, underscore."""
        docx_content = _create_empty_docx()
        valid_ids = ["abc", "ABC-123", "test_template", "a-b_c-1", "A" * 100]

        for tid in valid_ids:
            response = client.post(
                "/upload-template",
                files={"file": ("template.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"template_id": tid, "section": "cover"},
            )
            assert response.status_code == 201, f"Failed for template_id: {tid}"
