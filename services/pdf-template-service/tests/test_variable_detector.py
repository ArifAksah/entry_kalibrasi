"""Unit tests untuk variable_detector.py."""

import tempfile
from pathlib import Path

import pytest
from docx import Document

from app.services.variable_detector import detect_variables


@pytest.fixture
def tmp_dir():
    """Temporary directory untuk test files."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


def _create_docx_with_text(path: Path, paragraphs: list[str], table_cells: list[str] | None = None):
    """Helper: buat file .docx dengan paragraf dan opsional tabel."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)

    if table_cells:
        # Buat tabel 1 kolom dengan N baris
        table = doc.add_table(rows=len(table_cells), cols=1)
        for i, cell_text in enumerate(table_cells):
            table.rows[i].cells[0].text = cell_text

    doc.save(str(path))


class TestDetectVariables:
    """Tests untuk fungsi detect_variables."""

    def test_detects_simple_variables(self, tmp_dir):
        """Deteksi variabel sederhana {{ nama }}."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(docx_path, [
            "Nama alat: {{ nama_alat }}",
            "Merk: {{ merk }}",
            "Tipe: {{ tipe }}",
        ])

        variables, loops = detect_variables(docx_path)

        assert "nama_alat" in variables
        assert "merk" in variables
        assert "tipe" in variables
        assert len(variables) == 3
        assert loops == []

    def test_detects_loop_variables(self, tmp_dir):
        """Deteksi loop {% for item in collection %}."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(docx_path, [
            "{% for sensor in sensors %}",
            "Sensor: {{ sensor_nama }}",
            "{% endfor %}",
        ])

        variables, loops = detect_variables(docx_path)

        assert "sensor_nama" in variables
        assert "sensors" in loops

    def test_detects_variables_in_table_cells(self, tmp_dir):
        """Deteksi variabel di dalam tabel."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(
            docx_path,
            paragraphs=["Header dokumen"],
            table_cells=[
                "{{ nomor_sertifikat }}",
                "{{ tanggal_kalibrasi }}",
            ],
        )

        variables, loops = detect_variables(docx_path)

        assert "nomor_sertifikat" in variables
        assert "tanggal_kalibrasi" in variables

    def test_removes_duplicates(self, tmp_dir):
        """Variabel duplikat hanya muncul sekali."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(docx_path, [
            "{{ nama_alat }} - {{ merk }}",
            "Alat: {{ nama_alat }}",
            "Merk: {{ merk }}",
        ])

        variables, loops = detect_variables(docx_path)

        assert variables.count("nama_alat") == 1
        assert variables.count("merk") == 1
        assert len(variables) == 2

    def test_empty_document(self, tmp_dir):
        """Dokumen kosong mengembalikan list kosong."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(docx_path, ["Dokumen tanpa variabel"])

        variables, loops = detect_variables(docx_path)

        assert variables == []
        assert loops == []

    def test_multiple_loops(self, tmp_dir):
        """Deteksi multiple loop collections."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(docx_path, [
            "{% for sensor in sensors %}",
            "{% for hasil in hasil_kalibrasi %}",
            "{{ titik_ukur }}",
            "{% endfor %}",
            "{% endfor %}",
        ])

        variables, loops = detect_variables(docx_path)

        assert "sensors" in loops
        assert "hasil_kalibrasi" in loops
        assert "titik_ukur" in variables

    def test_variables_with_extra_spaces(self, tmp_dir):
        """Variabel dengan spasi ekstra tetap terdeteksi."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(docx_path, [
            "{{  nama_alat  }}",
            "{{nama_alat}}",
            "{{ merk }}",
        ])

        variables, loops = detect_variables(docx_path)

        # Semua variasi harus terdeteksi sebagai nama yang sama
        assert "nama_alat" in variables
        assert "merk" in variables

    def test_loop_with_whitespace_variations(self, tmp_dir):
        """Loop dengan variasi whitespace tetap terdeteksi."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(docx_path, [
            "{%  for item in items  %}",
            "{%- for sensor in sensors -%}",
            "{% for x in data %}",
        ])

        variables, loops = detect_variables(docx_path)

        assert "items" in loops
        assert "sensors" in loops
        assert "data" in loops

    def test_file_not_found_raises_error(self, tmp_dir):
        """File yang tidak ada menghasilkan error."""
        nonexistent = tmp_dir / "nonexistent.docx"

        with pytest.raises(Exception):
            detect_variables(nonexistent)

    def test_mixed_content(self, tmp_dir):
        """Dokumen dengan campuran variabel, loop, dan teks biasa."""
        docx_path = tmp_dir / "test.docx"
        _create_docx_with_text(
            docx_path,
            paragraphs=[
                "SERTIFIKAT KALIBRASI",
                "Nomor: {{ nomor_sertifikat }}",
                "Alat: {{ nama_alat }}",
                "{% for sensor in sensors %}",
                "Sensor: {{ sensor_nama }}",
                "{% endfor %}",
            ],
            table_cells=[
                "{{ merk }}",
                "{{ tipe }}",
            ],
        )

        variables, loops = detect_variables(docx_path)

        assert set(variables) == {"nomor_sertifikat", "nama_alat", "sensor_nama", "merk", "tipe"}
        assert loops == ["sensors"]
