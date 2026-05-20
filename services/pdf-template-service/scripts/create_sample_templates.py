"""
Script untuk membuat sample template .docx dengan variabel Jinja2.
Jalankan di dalam Docker container atau environment yang punya docxtpl.

Usage:
  docker exec pdf-template-service-pdf-template-service-1 python /app/scripts/create_sample_templates.py

Atau jalankan langsung jika punya docxtpl terinstall:
  python scripts/create_sample_templates.py
"""

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "sample-templates")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_cover_template():
    """Buat template cover sertifikat kalibrasi dengan variabel Jinja2."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header - BMKG
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BADAN METEOROLOGI KLIMATOLOGI DAN GEOFISIKA")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LABORATORIUM KALIBRASI BMKG")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SERTIFIKAT KALIBRASI")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CALIBRATION CERTIFICATE")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # Nomor Sertifikat
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Nomor / Number: ")
    run.font.size = Pt(11)
    run = p.add_run("{{ nomor_sertifikat }}")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Data Instrumen
    p = doc.add_paragraph()
    run = p.add_run("DATA INSTRUMEN / INSTRUMENT DATA")
    run.bold = True
    run.font.size = Pt(11)

    # Table for instrument data
    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    data_rows = [
        ("Nama Alat / Instrument", "{{ nama_alat }}"),
        ("Merk / Brand", "{{ merk }}"),
        ("Tipe / Type", "{{ tipe }}"),
        ("No. Seri / Serial Number", "{{ no_seri }}"),
        ("Kapasitas / Capacity", "{{ kapasitas }}"),
        ("Resolusi / Resolution", "{{ resolusi }}"),
        ("Satuan / Unit", "{{ unit }}"),
        ("Lain-lain / Others", "{{ lain_lain }}"),
    ]

    for i, (label, value) in enumerate(data_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()  # spacer

    # Data Pemilik
    p = doc.add_paragraph()
    run = p.add_run("DATA PEMILIK / OWNER DATA")
    run.bold = True
    run.font.size = Pt(11)

    table2 = doc.add_table(rows=3, cols=2)
    owner_rows = [
        ("Nama Stasiun / Station", "{{ nama_stasiun }}"),
        ("Alamat / Address", "{{ alamat_stasiun }}"),
        ("Pemilik / Owner", "{{ nama_pemilik }}"),
    ]

    for i, (label, value) in enumerate(owner_rows):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = value

    doc.add_paragraph()  # spacer

    # Data Kalibrasi
    p = doc.add_paragraph()
    run = p.add_run("DATA KALIBRASI / CALIBRATION DATA")
    run.bold = True
    run.font.size = Pt(11)

    table3 = doc.add_table(rows=7, cols=2)
    cal_rows = [
        ("No. Order / Order Number", "{{ no_order }}"),
        ("Tanggal Masuk / Received Date", "{{ tanggal_masuk }}"),
        ("Tanggal Kalibrasi / Calibration Date", "{{ tanggal_kalibrasi }}"),
        ("Tanggal Terbit / Issue Date", "{{ tanggal_terbit }}"),
        ("Metode / Method", "{{ metode_kalibrasi }}"),
        ("Suhu / Temperature", "{{ suhu }} °C"),
        ("Kelembaban / Humidity", "{{ kelembaban }} %"),
    ]

    for i, (label, value) in enumerate(cal_rows):
        table3.rows[i].cells[0].text = label
        table3.rows[i].cells[1].text = value

    doc.add_paragraph()  # spacer

    # Pengesahan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ nama_penandatangan }}")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("NIP. {{ nip_penandatangan }}")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("{{ jabatan_penandatangan }}")
    run.font.size = Pt(10)

    # Save
    output_path = os.path.join(OUTPUT_DIR, "fc-cover-template.docx")
    doc.save(output_path)
    print(f"✅ Cover template created: {output_path}")
    return output_path


def create_results_template():
    """Buat template hasil kalibrasi dengan loop Jinja2 untuk sensors."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header kecil
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HASIL KALIBRASI / CALIBRATION RESULTS")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("No. {{ nomor_sertifikat }}")
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Loop per sensor
    p = doc.add_paragraph()
    run = p.add_run("{% for sensor in sensors %}")
    run.font.size = Pt(1)  # hide the loop tag visually

    # Sensor info
    p = doc.add_paragraph()
    run = p.add_run("Sensor: {{ sensor.sensor_nama }} ({{ sensor.sensor_merk }} {{ sensor.sensor_tipe }})")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("No. Seri: {{ sensor.sensor_no_seri }}")
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Table header for results
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["No.", "Titik Ukur", "Pembacaan", "Koreksi", "Ketidakpastian"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Data row (loop)
    table.rows[1].cells[0].text = "{% for row in sensor.hasil_kalibrasi %}{{ row.no_urut }}"
    table.rows[1].cells[1].text = "{{ row.titik_ukur }}"
    table.rows[1].cells[2].text = "{{ row.pembacaan }}"
    table.rows[1].cells[3].text = "{{ row.koreksi }}"
    table.rows[1].cells[4].text = "{{ row.ketidakpastian }}{% endfor %}"

    doc.add_paragraph()  # spacer

    # Catatan
    p = doc.add_paragraph()
    run = p.add_run("Catatan: {{ catatan }}")
    run.font.size = Pt(10)

    # End sensor loop
    p = doc.add_paragraph()
    run = p.add_run("{% endfor %}")
    run.font.size = Pt(1)  # hide

    doc.add_paragraph()  # spacer

    # Teknisi
    p = doc.add_paragraph()
    run = p.add_run("Teknisi Pelaksana: {{ nama_teknisi }} (NIP: {{ nip_teknisi }})")
    run.font.size = Pt(10)

    # Save
    output_path = os.path.join(OUTPUT_DIR, "fc-results-template.docx")
    doc.save(output_path)
    print(f"✅ Results template created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Membuat sample templates...")
    create_cover_template()
    create_results_template()
    print(f"\n📁 Templates tersimpan di: {OUTPUT_DIR}")
    print("\nGunakan file-file ini untuk test upload di halaman admin/templates/[id]/word-upload")
