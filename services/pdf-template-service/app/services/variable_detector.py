"""Deteksi variabel Jinja2 ({{ ... }} dan {% for ... %}) dalam file .docx."""

import re
from pathlib import Path

from docx import Document


# Regex untuk mendeteksi {{ variable_name }} — menangkap nama variabel saja
_VARIABLE_PATTERN = re.compile(r"\{\{\s*(\w+)\s*\}\}")

# Regex untuk mendeteksi {% for item in collection %} — menangkap nama collection
_LOOP_PATTERN = re.compile(r"\{%[-\s]*for\s+\w+\s+in\s+(\w+)\s*[-]?%\}")


def _extract_all_text(doc: Document) -> str:
    """Ekstrak semua teks dari paragraf dan tabel dalam dokumen .docx.

    docxtpl menyimpan variabel Jinja2 di dalam paragraf dan tabel cell.
    Kita perlu membaca semua teks untuk menemukan pola {{ }} dan {% %}.
    """
    texts: list[str] = []

    # Teks dari semua paragraf
    for paragraph in doc.paragraphs:
        texts.append(paragraph.text)

    # Teks dari semua tabel (setiap cell bisa mengandung variabel)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    texts.append(paragraph.text)

    return "\n".join(texts)


def detect_variables(template_path: Path) -> tuple[list[str], list[str]]:
    """Deteksi semua Template_Variable dan Loop_Variable dalam file .docx.

    Membuka file .docx menggunakan python-docx, membaca semua teks dari
    paragraf dan tabel, lalu mencari pola:
    - {{ variable_name }} → variabel template
    - {% for item in collection %} → loop/collection

    Args:
        template_path: Path ke file .docx yang akan dianalisis.

    Returns:
        Tuple berisi:
        - variables: list nama variabel unik (tanpa {{ }})
        - loops: list nama collection unik (tanpa {% %})

    Raises:
        FileNotFoundError: Jika file tidak ditemukan.
        Exception: Jika file bukan format .docx yang valid.
    """
    doc = Document(str(template_path))
    full_text = _extract_all_text(doc)

    # Cari semua variabel {{ ... }} — deduplicate dengan dict.fromkeys (preserve order)
    variable_matches = _VARIABLE_PATTERN.findall(full_text)
    variables = list(dict.fromkeys(variable_matches))

    # Cari semua loop {% for ... in collection %} — deduplicate
    loop_matches = _LOOP_PATTERN.findall(full_text)
    loops = list(dict.fromkeys(loop_matches))

    return variables, loops
